from docx import Document
from docx.table import Table
from docx.text.run import Run
from pandas import DataFrame


def replace_placeholder_paragraph(paragraph: Run, replacements):
    """尝试替换整个段落为指定的占位符值"""
    text = paragraph.text.strip()
    for old_text, new_text in replacements.items():
        if text == old_text:
            paragraph.text = str(new_text)
            return True
    return False


def replace_placeholders_in_table(table: Table, replacements):
    """在给定的表格中替换占位符"""
    for row in table.rows:
        for cell in row.cells:
            # 如果单元格包含段落
            if cell.paragraphs:
                for paragraph in cell.paragraphs:
                    # 尝试替换整个段落
                    if not replace_placeholder_paragraph(paragraph, replacements):
                        # 如果没有找到匹配的占位符，则遍历段落中的run
                        for run in paragraph.runs:
                            text = run.text
                            # print(f"当前run:{text}")
                            if not text.strip():
                                continue

                            for old_text, new_text in replacements.items():
                                # print(f"占位符KEY:{old_text},Value:{new_text}")
                                if old_text in text:
                                    # 只有当整个run的文本是占位符时才替换
                                    if text.strip() == old_text:
                                        run.text = str(new_text)


def replace_placeholders_in_docx(doc: Document, replacements):
    """遍历文档中的所有表格，并替换其中的占位符"""
    for table in doc.tables:
        replace_placeholders_in_table(table, replacements)


def ret_fault_detection(end_time, begin_time, replacements, err_msg, df: DataFrame):
    end_time = end_time or (begin_time + 5)
    condition5 = df['timestamps'] >= begin_time
    condition6 = df['timestamps'] <= end_time
    draw_detection_figure_df = df[condition5 & condition6]
    return err_msg, replacements, draw_detection_figure_df
